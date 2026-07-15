#!/usr/bin/env python3
"""Build the Chord Markov Performer v4 user manual.

Design preset: compact_reference_guide
First-page template: editorial_cover

The Markdown file is the maintainable source. This script creates a polished
Word edition with a cover, running furniture, fixed-width tables, real list
numbering, callouts, and a schematic of the compact Max for Live panel.
"""

from __future__ import annotations

import argparse
import re
import tempfile
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


HERE = Path(__file__).resolve().parent
SOURCE = HERE / "USER_MANUAL_V4.md"
DEFAULT_OUTPUT = HERE / "Chord_Markov_Performer_v4_User_Manual.docx"

FONT_BODY = "Calibri"
FONT_MONO = "Consolas"

NAVY = "17324D"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
PALE_BLUE = "F4F8FC"
TEXT = "24313D"
MUTED = "5E6B76"
LINE = "B8C4CF"
PALE_GRAY = "F3F5F7"
WHITE = "FFFFFF"
AMBER = "E59B3A"
PALE_AMBER = "FFF4DE"

PAGE_W = Inches(8.5)
PAGE_H = Inches(11)
MARGIN = Inches(1)
HEADER_FOOTER = Inches(0.492)
CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), LINE)

    for row in table.rows:
        row_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        row_pr.append(cant_split)
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    header_pr = table.rows[0]._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    header_pr.append(repeat)


def set_repeatable_row(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_paragraph_border(paragraph, *, side: str, color: str, size: int = 12) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    border = p_bdr.find(qn(f"w:{side}"))
    if border is None:
        border = OxmlElement(f"w:{side}")
        p_bdr.append(border)
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(size))
    border.set(qn("w:space"), "4")
    border.set(qn("w:color"), color)


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_keep(paragraph, *, next_: bool = False, together: bool = False) -> None:
    paragraph.paragraph_format.keep_with_next = next_
    paragraph.paragraph_format.keep_together = together


def set_font(run, name: str, size: float, color: str | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def restart_page_number(section, start: int = 1) -> None:
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    pg_num.set(qn("w:start"), str(start))


def set_section_page(section) -> None:
    section.page_width = PAGE_W
    section.page_height = PAGE_H
    section.top_margin = MARGIN
    section.bottom_margin = MARGIN
    section.left_margin = MARGIN
    section.right_margin = MARGIN
    section.header_distance = HEADER_FOOTER
    section.footer_distance = HEADER_FOOTER


def configure_styles(doc: Document) -> None:
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = FONT_BODY
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    fmt = normal.paragraph_format
    fmt.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(6)
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = 1.25
    fmt.widow_control = True

    heading_specs = {
        "Title": (30, NAVY, 0, 12),
        "Subtitle": (17, MUTED, 0, 18),
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for style_name, (size, color, before, after) in heading_specs.items():
        style = styles[style_name]
        style.font.name = FONT_BODY
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.widow_control = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = FONT_BODY
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    caption = styles["Caption"]
    caption.font.name = FONT_BODY
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(10)
    caption.paragraph_format.keep_with_next = False


def set_furniture_borders(table, *, bottom_color: str | None = None) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        if edge == "bottom" and bottom_color:
            element.set(qn("w:val"), "single")
            element.set(qn("w:sz"), "8")
            element.set(qn("w:color"), bottom_color)
        else:
            element.set(qn("w:val"), "nil")


def add_furniture_table(container, *, header: bool) -> None:
    table = container.add_table(rows=1, cols=2, width=Inches(6.5))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    widths = [4680, 4680]
    for idx, cell in enumerate(table.rows[0].cells):
        set_cell_width(cell, widths[idx])
        set_cell_margins(cell, top=0, bottom=20, start=0, end=0)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_furniture_borders(table, bottom_color=LIGHT_BLUE if header else None)

    left_p = table.cell(0, 0).paragraphs[0]
    right_p = table.cell(0, 1).paragraphs[0]
    left_p.paragraph_format.space_after = Pt(0)
    right_p.paragraph_format.space_after = Pt(0)
    right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    if header:
        left = left_p.add_run("CHORD MARKOV PERFORMER v4")
        set_font(left, FONT_BODY, 8.5, BLUE)
        left.bold = True
        right = right_p.add_run("USER MANUAL  •  4.0.0-dev")
        set_font(right, FONT_BODY, 8.5, MUTED)
    else:
        left = left_p.add_run("FinalMaxUPF  •  Theory Prototype")
        set_font(left, FONT_BODY, 8.5, MUTED)
        page = right_p.add_run("Page ")
        set_font(page, FONT_BODY, 8.5, MUTED)
        add_field(right_p, "PAGE")

    # Word requires a trailing paragraph in a header/footer. Put the table
    # first and collapse that required paragraph to avoid visible extra space.
    container._element.remove(table._element)
    container._element.insert(0, table._element)
    trailing = container.paragraphs[-1]
    trailing.paragraph_format.space_before = Pt(0)
    trailing.paragraph_format.space_after = Pt(0)
    trailing.paragraph_format.line_spacing = Pt(1)
    if not trailing.runs:
        run = trailing.add_run("")
        set_font(run, FONT_BODY, 1, MUTED)


def add_running_header_footer(section) -> None:
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    add_furniture_table(section.header, header=True)
    add_furniture_table(section.footer, header=False)


def add_cover(doc: Document) -> None:
    section = doc.sections[0]
    set_section_page(section)
    section.different_first_page_header_footer = True

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(52)
    set_paragraph_border(p, side="top", color=BLUE, size=24)
    run = p.add_run("FINALMAXUPF  /  THEORY PROTOTYPE")
    set_font(run, FONT_BODY, 9, BLUE)
    run.bold = True

    p = doc.add_paragraph()
    p.style = doc.styles["Title"]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Chord Markov")
    set_font(run, FONT_BODY, 30, NAVY)
    run.bold = True

    p = doc.add_paragraph()
    p.style = doc.styles["Title"]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("Performer v4")
    set_font(run, FONT_BODY, 34, BLUE)
    run.bold = True

    p = doc.add_paragraph()
    p.style = doc.styles["Subtitle"]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("User Manual")
    set_font(run, FONT_BODY, 17, MUTED)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(26)
    run = p.add_run("Models, harmony, rhythm, performance, and recovery")
    set_font(run, FONT_BODY, 11, DARK_BLUE)

    meta = doc.add_table(rows=3, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.autofit = False
    cover_widths = [2200, 5000]
    labels = [
        ("DEVICE", "Chord Markov Performer v4.amxd"),
        ("PLATFORM", "Ableton Live + Max for Live"),
        ("REVISION", "4.0.0-dev  •  13 July 2026"),
    ]
    for row, (label, value) in zip(meta.rows, labels):
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, cover_widths[idx])
            set_cell_margins(cell, top=90, bottom=90, start=120, end=120)
            set_cell_shading(cell, PALE_BLUE if idx == 0 else WHITE)
        lp = row.cells[0].paragraphs[0]
        lr = lp.add_run(label)
        set_font(lr, FONT_BODY, 8.5, BLUE)
        lr.bold = True
        vp = row.cells[1].paragraphs[0]
        vr = vp.add_run(value)
        set_font(vr, FONT_BODY, 10.5, TEXT)
    set_table_geometry(meta, cover_widths)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(26)
    p.paragraph_format.space_after = Pt(4)
    set_paragraph_shading(p, PALE_AMBER)
    set_paragraph_border(p, side="left", color=AMBER, size=22)
    r = p.add_run("START WITH A FRESH DEVICE")
    set_font(r, FONT_BODY, 10, NAVY)
    r.bold = True
    p.add_run("\n")
    r = p.add_run(
        "Remove every loaded v4 instance, then drag the exact current AMXD into the "
        "Set again. Use one instance only; it owns local reply port 9101."
    )
    set_font(r, FONT_BODY, 10, TEXT)
    p.paragraph_format.left_indent = Inches(0.14)
    p.paragraph_format.right_indent = Inches(0.14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    r = p.add_run(
        "This edition documents the repaired full-chord MIDI path, theory-aware "
        "Complexity, five generators, seven Feel patterns, and self-launching backend."
    )
    set_font(r, FONT_BODY, 9, MUTED)


def create_numbering(doc: Document) -> tuple[int, int]:
    numbering = doc.part.numbering_part.element
    existing_abs = [
        int(n.get(qn("w:abstractNumId")))
        for n in numbering.findall(qn("w:abstractNum"))
        if n.get(qn("w:abstractNumId")) is not None
    ]
    next_abs = max(existing_abs, default=-1) + 1

    def abstract(abstract_id: int, fmt: str, text: str, font: str | None = None) -> None:
        abs_num = OxmlElement("w:abstractNum")
        abs_num.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abs_num.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        suffix = OxmlElement("w:suff")
        suffix.set(qn("w:val"), "tab")
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "540")
        tabs.append(tab)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "540")
        ind.set(qn("w:hanging"), "270")
        p_pr.extend([tabs, ind])
        lvl.extend([start, num_fmt, lvl_text, suffix, p_pr])
        if font:
            r_pr = OxmlElement("w:rPr")
            fonts = OxmlElement("w:rFonts")
            fonts.set(qn("w:ascii"), font)
            fonts.set(qn("w:hAnsi"), font)
            r_pr.append(fonts)
            lvl.append(r_pr)
        abs_num.append(lvl)
        numbering.append(abs_num)

    decimal_abs = next_abs
    bullet_abs = next_abs + 1
    abstract(decimal_abs, "decimal", "%1.")
    abstract(bullet_abs, "bullet", "•", "Symbol")
    return decimal_abs, bullet_abs


def new_num_id(doc: Document, abstract_id: int, start: int = 1) -> int:
    numbering = doc.part.numbering_part.element
    existing = [
        int(n.get(qn("w:numId")))
        for n in numbering.findall(qn("w:num"))
        if n.get(qn("w:numId")) is not None
    ]
    num_id = max(existing, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), str(abstract_id))
    num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), str(start))
    override.append(start_override)
    num.append(override)
    numbering.append(num)
    return num_id


def apply_num(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    n_id = OxmlElement("w:numId")
    n_id.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, n_id])


def add_inline(paragraph, text: str, *, base_size: float = 11, base_color: str = TEXT) -> None:
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        bold = part.startswith("**") and part.endswith("**")
        clean = part[2:-2] if bold else part
        clean = clean.replace("  ", " ")
        run = paragraph.add_run(clean)
        set_font(run, FONT_BODY, base_size, base_color)
        run.bold = bold


def add_body_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    add_inline(p, text)


def add_callout(doc: Document, text: str) -> None:
    important = text.upper().startswith("IMPORTANT")
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.16)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(9)
    p.paragraph_format.keep_together = True
    set_paragraph_shading(p, PALE_AMBER if important else PALE_BLUE)
    set_paragraph_border(p, side="left", color=AMBER if important else BLUE, size=20)
    add_inline(p, text, base_size=10.25)


def add_code(doc: Document, lines: list[str]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_together = True
    set_paragraph_shading(p, PALE_GRAY)
    set_paragraph_border(p, side="left", color=BLUE, size=12)
    for idx, line in enumerate(lines):
        if idx:
            p.add_run("\n")
        run = p.add_run(line)
        set_font(run, FONT_MONO, 8.8, DARK_BLUE)


def table_widths(rows: list[list[str]]) -> list[int]:
    cols = len(rows[0])
    header = [re.sub(r"\*\*", "", cell).lower() for cell in rows[0]]
    if cols == 2:
        ratios = [0.27, 0.73]
    elif cols == 3:
        ratios = [0.22, 0.28, 0.50]
        if "likely cause" in header:
            ratios = [0.27, 0.30, 0.43]
        elif "template" in header:
            ratios = [0.18, 0.26, 0.56]
    elif cols == 4:
        ratios = [0.16, 0.20, 0.30, 0.34]
        if "model" in header:
            ratios = [0.13, 0.31, 0.20, 0.36]
        elif "tier" in header:
            ratios = [0.14, 0.13, 0.39, 0.34]
    else:
        lengths = []
        for idx in range(cols):
            maximum = max(len(re.sub(r"\*\*", "", row[idx])) for row in rows)
            lengths.append(max(8, min(maximum, 40)))
        total = sum(lengths)
        ratios = [value / total for value in lengths]
    widths = [int(CONTENT_DXA * ratio) for ratio in ratios]
    widths[-1] += CONTENT_DXA - sum(widths)
    return widths


def add_table(doc: Document, rows: list[list[str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    widths = table_widths(rows)
    for r_idx, values in enumerate(rows):
        row = table.rows[r_idx]
        for c_idx, value in enumerate(values):
            cell = row.cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(1 if r_idx else 0)
            p.paragraph_format.line_spacing = 1.05
            if r_idx == 0:
                set_cell_shading(cell, LIGHT_BLUE)
                add_inline(p, value, base_size=9.2, base_color=NAVY)
                for run in p.runs:
                    run.bold = True
            else:
                if r_idx % 2 == 0:
                    set_cell_shading(cell, "FAFBFC")
                add_inline(p, value, base_size=9.2)
        row.height = None
    set_table_geometry(table, widths)
    set_repeatable_row(table.rows[0])
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(0)
    after.paragraph_format.space_before = Pt(0)
    after.paragraph_format.line_spacing = 0.2


def find_font() -> str | None:
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf",
        "/Library/Fonts/Arial.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    ]
    return next((path for path in candidates if Path(path).is_file()), None)


def draw_panel_map(path: Path) -> None:
    width, height = 1800, 500
    image = Image.new("RGB", (width, height), "#F5F7FA")
    draw = ImageDraw.Draw(image)
    font_path = find_font()
    title_font = ImageFont.truetype(font_path, 34) if font_path else ImageFont.load_default()
    label_font = ImageFont.truetype(font_path, 25) if font_path else ImageFont.load_default()
    small_font = ImageFont.truetype(font_path, 20) if font_path else ImageFont.load_default()

    draw.rounded_rectangle((12, 12, width - 12, height - 12), radius=24, fill="#141A21", outline="#2E74B5", width=5)
    zones = [
        (30, 32, 210, 390, "#26364A", "TRANSPORT", ["Play • Reroll", "Hold • Free/Sync", "Voices • BPM"]),
        (225, 32, 470, 390, "#1F4D78", "MODEL", ["Markov • RNN", "LSTM • n-gram", "Phrase", "session • capture • v4"]),
        (485, 32, 810, 390, "#2E5B72", "PHRASE + THEORY", ["Bars • Mode", "Cadence", "Complexity"]),
        (825, 32, 1500, 390, "#334554", "SELECTION + SOUND", ["Seed • Key • maj/min", "Spice • Rhythm • Voicing", "backend • Relink"]),
        (1515, 32, 1770, 390, "#5B466B", "FEEL", ["Straight • Push", "Tresillo • Clave", "Upbeats • Triplet"]),
    ]
    for x1, y1, x2, y2, fill, title, lines in zones:
        draw.rounded_rectangle((x1, y1, x2, y2), radius=18, fill=fill, outline="#6F8294", width=2)
        draw.text((x1 + 18, y1 + 18), title, font=label_font, fill="#FFFFFF")
        y = y1 + 84
        for line in lines:
            draw.text((x1 + 18, y), line, font=small_font, fill="#E8EEF5")
            y += 54
    draw.rounded_rectangle((30, 410, 1770, 468), radius=12, fill="#0C1117", outline="#506477", width=2)
    draw.text((52, 421), "OUTPUT  —  CURRENT NORMALIZED CHORD SYMBOL", font=title_font, fill="#A8D8FF")
    image.save(path, dpi=(180, 180))


def add_panel_map(doc: Document, image_path: Path) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(6.5))
    caption = doc.add_paragraph(style="Caption")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.add_run("Figure 1. Current v4 panel zones, arranged left to right.")


def parse_markdown(doc: Document, lines: list[str], panel_image: Path) -> None:
    decimal_abs, bullet_abs = create_numbering(doc)
    bullet_num = new_num_id(doc, bullet_abs)
    active_number_num: int | None = None
    active_list_kind: str | None = None
    paragraph_buffer: list[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph_buffer
        if paragraph_buffer:
            add_body_paragraph(doc, " ".join(part.strip() for part in paragraph_buffer))
            paragraph_buffer = []

    i = 0
    while i < len(lines):
        raw = lines[i].rstrip()
        stripped = raw.strip()

        if stripped == "<!-- PAGE BREAK -->":
            flush_paragraph()
            doc.add_page_break()
            active_list_kind = None
            active_number_num = None
            i += 1
            continue

        if not stripped:
            flush_paragraph()
            active_list_kind = None
            active_number_num = None
            i += 1
            continue

        if stripped.startswith("#"):
            flush_paragraph()
            match = re.match(r"^(#{1,3})\s+(.*)$", stripped)
            if match:
                level = len(match.group(1))
                text = match.group(2)
                p = doc.add_paragraph(style=f"Heading {level}")
                add_inline(
                    p,
                    text,
                    base_size={1: 16, 2: 13, 3: 12}[level],
                    base_color={1: BLUE, 2: BLUE, 3: DARK_BLUE}[level],
                )
                for run in p.runs:
                    run.bold = True
                if text == "4. Reading the panel":
                    add_panel_map(doc, panel_image)
            i += 1
            continue

        if stripped.startswith(">"):
            flush_paragraph()
            quote_lines: list[str] = []
            while i < len(lines) and lines[i].lstrip().startswith(">"):
                content = lines[i].lstrip()[1:].strip()
                if content:
                    quote_lines.append(content)
                i += 1
            add_callout(doc, " ".join(quote_lines))
            continue

        if raw.startswith("    "):
            flush_paragraph()
            code_lines: list[str] = []
            while i < len(lines) and (lines[i].startswith("    ") or not lines[i].strip()):
                if lines[i].startswith("    "):
                    code_lines.append(lines[i][4:].rstrip())
                elif code_lines and i + 1 < len(lines) and lines[i + 1].startswith("    "):
                    code_lines.append("")
                i += 1
            add_code(doc, code_lines)
            continue

        if stripped.startswith("|") and i + 1 < len(lines):
            delimiter = lines[i + 1].strip()
            if delimiter.startswith("|") and re.fullmatch(r"[|:\-\s]+", delimiter):
                flush_paragraph()
                table_lines = [stripped]
                i += 2
                while i < len(lines) and lines[i].strip().startswith("|"):
                    table_lines.append(lines[i].strip())
                    i += 1
                rows = []
                for table_line in table_lines:
                    cells = [cell.strip() for cell in table_line.strip("|").split("|")]
                    rows.append(cells)
                if not rows or not all(len(row) == len(rows[0]) for row in rows):
                    raise RuntimeError(f"Malformed Markdown table near: {table_lines[0]}")
                add_table(doc, rows)
                continue

        number_match = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if number_match:
            flush_paragraph()
            if active_list_kind != "number":
                active_number_num = new_num_id(doc, decimal_abs, int(number_match.group(1)))
                active_list_kind = "number"
            p = doc.add_paragraph(style="List Number")
            apply_num(p, active_number_num)
            add_inline(p, number_match.group(2))
            i += 1
            continue

        if stripped.startswith("- "):
            flush_paragraph()
            active_list_kind = "bullet"
            p = doc.add_paragraph(style="List Bullet")
            apply_num(p, bullet_num)
            add_inline(p, stripped[2:])
            i += 1
            continue

        paragraph_buffer.append(stripped)
        i += 1

    flush_paragraph()


def build(source: Path, output: Path) -> None:
    text = source.read_text(encoding="utf-8")
    lines = text.splitlines()
    try:
        first_break = lines.index("<!-- PAGE BREAK -->")
    except ValueError as exc:
        raise RuntimeError("manual source is missing its first page-break marker") from exc
    body_lines = lines[first_break + 1 :]

    doc = Document()
    configure_styles(doc)
    doc.core_properties.title = "Chord Markov Performer v4 User Manual"
    doc.core_properties.subject = "Ableton Live and Max for Live operating manual"
    doc.core_properties.author = "FinalMaxUPF"
    doc.core_properties.keywords = "Max for Live, chord generation, harmony, rhythm, user manual"
    doc.core_properties.comments = (
        "Built from USER_MANUAL_V4.md using compact_reference_guide and editorial_cover."
    )

    add_cover(doc)
    body_section = doc.add_section(WD_SECTION.NEW_PAGE)
    set_section_page(body_section)
    body_section.different_first_page_header_footer = False
    restart_page_number(body_section, 1)
    add_running_header_footer(body_section)

    with tempfile.TemporaryDirectory(prefix="v4-user-manual-") as temp_dir:
        panel_image = Path(temp_dir) / "panel-map.png"
        draw_panel_map(panel_image)
        parse_markdown(doc, body_lines, panel_image)
        output.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output)


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--source", type=Path, default=SOURCE)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    args = parser.parse_args()
    build(args.source.resolve(), args.output.resolve())
    print(f"wrote {args.output.resolve()}")


if __name__ == "__main__":
    main()
